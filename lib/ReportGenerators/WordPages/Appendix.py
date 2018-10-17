from ...DataAbstractions import *
from ..WordGenerator import Organizer

import docx
import logging
from lxml import etree
import urllib.request
import zipfile
import io

#
# Base class - ReportGenerator
#
class AppendixPage:
  #
  # Name of the page
  #
  @staticmethod
  def name ():
    return 'Appendix'

  #
  # Constructor
  #
  def __init__ (self):
    self._weaknesses = set () 
    self._catalog_url = "http://cwe.mitre.org/data/xml/cwec_v2.4.xml.zip"

  #
  # Initialize the page - Generic
  #
  def init (self, organizer, wrong_checker_is_fp):
    self.init ()

  #
  # Initalize the page - Appendix-specific
  #
  def init (self):
    pass

  #
  # Visit datapoint
  #
  def visit (self, datapoint):
    self._weaknesses.add (datapoint.weakness)

  #
  # fini
  #
  def fini (self, document):
    catalog = self.get_catalog ()
    document.add_heading ('Weakness Details')

    for weakness in sorted (self._weaknesses):
      # Use xpath to find the weakness in the catalog, catalog has the weakness number
      # as the ID whereas we have a 'CWE' prefix
      info = catalog.xpath ('Weaknesses/Weakness[@ID=%s]' % weakness[3:])
      if len (info) == 0:
        logging.warning ('Unable to find Weakness [%s] in catalog' % weakness[3:])
        continue

      self.write_appendix_entry (info[0], document)

  #
  # Get the catalog from NIST
  #
  def get_catalog (self):
    # NIST stores the catalog in a zip file, so we must download it,
    # extract the XML document from the zip, then load it up into
    # an lxml.etree object
    logging.info ('Downloading catalog from [%s]' % self._catalog_url)
    response = urllib.request.urlopen (self._catalog_url).add_results ()
    zip_f = zipfile.ZipFile (io.BytesIO (response))
    xml_f = zip_f.open ('cwec_v2.4.xml')
    return etree.fromstring (xml_f.add_results ())

  #
  # Write a single appendix entry
  #
  def write_appendix_entry (self, entry, document):
    # @TODO: See if using the XSD will make this easier
    (cwe, name, abstraction, status) = entry.values ()
    target = 'CWE' + cwe

    # Write subsection and header
    document.add_heading ('%s - %s' % (target, name), level=4)

    # Write description
    desc = entry.findall('Description/Description_Summary')
    for i in desc:
      document.add_paragraph ('Description:\n%s\n' % self.remove_formatting (i.text))

    # Write extended description
    extdesc = entry.findall ('Description/Extended_Description/Text')
    if len (extdesc) != 0:
      p = document.add_paragraph ('Extended Description:\n')

      for i in extdesc:
        p.add_run ('%s' % self.remove_formatting (i.text))

    # Write platforms
    platforms = entry.findall ('Applicable_Platforms/Languages/Language')
    if len (platforms) != 0:
      document.add_paragraph ('Applicable Platforms:')

      for i in platforms:
        document.add_paragraph ('%s' % i.get ('Language_Name'), style='ListBullet')

    # Write consequences
    consequences = entry.findall ('Common_Consequences/Common_Consequence')
    if len (consequences) != 0:
      p = document.add_paragraph ('Common Consequences:')

      for consequence in consequences:
        scopes = consequence.findall ('Consequence_Scope')
        impacts = consequence.findall ('Consequence_Technical_Impact')
        notes = consequence.findall ('Consequence_Note/Text')

        if len (scopes) != 0:
          p.add_run ('\nConsequence Scope: %s\n' % ', '.join ([self.remove_formatting (s.text) for s in scopes]))

        if len (impacts) != 0:
          p.add_run ('Consequence Technical Impact: %s\n' % ', '.join ([self.remove_formatting (i.text) for i in impacts]))

        if len (notes) != 0:
          p.add_run ('Notes: %s\n' % self.remove_formatting (notes[0].text))

    # Write footer
    document.add_paragraph ('For more information please see: http://cwe.mitre.org/data/definitions/%s.html' % cwe)

  #
  # Remove formatting from the provided string
  #
  def remove_formatting (self, src):
    return src.replace ('\n', ' ').replace ('\t', '').replace ('_', ' ')
