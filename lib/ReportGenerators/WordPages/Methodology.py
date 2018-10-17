from ...DataAbstractions import *
from ... import Utilities
from ..WordGenerator import Organizer

import logging
import docx

#
# Base class - ReportGenerator
#
class MethodologyPage:
  #
  # Name of the page
  #
  @staticmethod
  def name ():
    return 'Methodology'

  #
  # Constructor
  #
  def __init__ (self):
    self._tool = None
    self._import_suite = None
    self._tool_args = None
    self._import_suite_args = None
    self._analyze_string = None
    self._wrong_checker_is_fp = None

  #
  # Initialize the page
  #
  def init (self, organizer, wrong_checker_is_fp):
    for tool in Utilities.get_tools ():
      if tool.name () == organizer.build_rs.GetSource ():
        self._tool = tool ()

    for import_suite in Utilities.get_importsuites ():
      if import_suite.name () == organizer.truth_rs.GetSource ():
        self._import_suite = import_suite ()

    self._tool_args = organizer.build_rs.GetArgs ()
    self._import_suite_args = organizer.truth_rs.GetArgs ()
    self._wrong_checker_is_fp = wrong_checker_is_fp

  #
  # Visit datapoint
  #
  def visit (self, datapoint):
    pass

  #
  # fini
  #
  def fini (self, document):
    document.add_heading ('Methodology')

    document.add_heading ('ImportSuite', level=4)
    self._import_suite.methodology (self._import_suite_args, document)

    document.add_heading ('Tool', level=4)
    self._tool.methodology (self._tool_args, document)

    document.add_heading ('Analysis', level=4)
    self.write_analysis (document)

    document.add_heading ('Granularity', level=4)
    self.write_granularity (document)

    document.add_page_break ()

  def write_analysis (self, document):
    document.add_paragraph ('''
After importing the Import and Build results, an analysis is executed to
identify true positives (TP), false positives (FP) and false negatives (FN).
The analysis uses the following criteria:
''')
    document.add_paragraph ('TP: The SCA tool found a Bug with the correct checker for the weakness in a file.', style='ListBullet')
    document.add_paragraph ('FN: The SCA tool found less TPs than expected in a file, the difference between the expected number of TPs and those found are FNs.', style='ListBullet')
    if self._wrong_checker_is_fp:
      document.add_paragraph ('FP: The SCA tool found more TPs than expected in a file, the difference between the number of TPs found and number expected are FPs. Additionally, any Bugs found with the wrong checker for the weakness are considered FPs.', style='ListBullet')
    else:
      document.add_paragraph ('FP: The SCA tool found more TPs than expected in a file, the difference between the number of TPs found and number expected are FPs.', style='ListBullet')

    document.add_paragraph ('To apply this critera, the following steps are taken:\n')

    document.add_paragraph ('All Flaws from the import file are grouped by granularity.', style='ListBullet')
    document.add_paragraph ('All Bugs identified by the tool from the build file are grouped by granularity.', style='ListBullet')
    document.add_paragraph ('Any Bugs reported by the tool on lines where Flaws were not expected are skipped.', style='ListBullet')
    document.add_paragraph ('Any Incidental Flaws which were correctly identified by the tool are skipped.', style='ListBullet')
    document.add_paragraph ('Fixes are removed from the expected number of TPs.', style='ListBullet')
    document.add_paragraph ('The above criteria for TP, FN, and FP is applied.', style='ListBullet')

  def write_granularity (self, document):
    document.add_paragraph ('''
This document includes the following granularities:''')

    document.add_paragraph ('Filename: Flaws and Bugs are grouped by filename', style='ListBullet')
    document.add_paragraph ('Function: Flaws and Bugs are grouped by filename and function', style='ListBullet')
    document.add_paragraph ('Line: Flaws and Bugs are grouped by filename, function, and line', style='ListBullet')

    document.add_paragraph ('''
These granularities can be used to show different behavioral aspects of the SCA tool.  For example, the Filename granularity cannot distinguish between 'good' and 'bad' functions
so the TP result will be higher than it is for other granularities.  However, it will also have the highest FP results because Bugs reported by the tool are not filtered to only the functions or lines of interest.''')
